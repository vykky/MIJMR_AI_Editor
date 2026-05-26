import streamlit as st
import pdfplumber
import google.generativeai as genai
import time
from docx import Document
import io

# 1. Page Configuration
st.set_page_config(page_title="MIJMR Editorial Board Desk", page_icon="🔖", layout="wide")

st.title("🔖 MIJMR - Editorial Board Evaluation Desk")
st.markdown("Upload the submitted research paper (PDF) for official board evaluation and screening.")

# 2. Fetch API Key from Streamlit Secrets
try:
    api_key = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("⚠️ Error: Gemini API Key is missing! Please configure it in Streamlit Secrets.")
    st.stop()

# 3. Function to Extract Text from PDF
def extract_text_from_pdf(pdf_file):
    text = ""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text
    except Exception as e:
        return f"Error extracting text: {e}"

# 4. Function to Analyze Text using Gemini 3.5 Flash (Multilingual)
def analyze_paper_with_gemini(text, api_key, standard):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-3.5-flash')
    
    if standard == "High Standard (International/Foreign Level)":
        internal_rules = "CRITERIA: Evaluate strictly based on top-tier International benchmarks. Be highly rigorous."
    else:
        internal_rules = """
        CRITERIA: Evaluate based on local university standard (Arts & Science level). 
        - If the paper is moderately okay (50% to 70% satisfactory), your default decision MUST be APPROVED / PROCEED TO PUBLISH: YES. 
        - If extremely poor (below 50% logic), provide a strict 'REQUIRED ACTIONABLE CORRECTIONS LIST'.
        STRICT RULE: NEVER use any robotic or AI buzzwords. Write exactly like a human Senior Professor.
        """

    prompt = f"""
    You are a Senior Professor and Chief Screening Editor for MIJMR. 
    Analyze the following research paper text and generate the official board report.
    
    INTERNAL EVALUATION RULES:
    {internal_rules}
    
    CRITICAL MULTILINGUAL RULE: 
    First, detect the primary language of the submitted research paper. 
    You MUST write the ENTIRE evaluation report (including ALL headings, observations, ratings, and final verdicts) in the EXACT SAME LANGUAGE as the uploaded paper. 
    (e.g., If the paper is in Tamil, write the report entirely in highly professional academic Tamil. If English, use formal English).
    
    Provide the output EXACTLY in the following format (Translate the headings to the detected language appropriately). Start directly with the title.
    
    **(Translate this Title): MIJMR Editorial Board - Official Review Statement**
    
    **1. (Translate): Core Observations & Evaluation:**
    (Write 1-2 natural paragraphs evaluating the paper's concept and methodology in the detected language).
    
    **2. (Translate): Document & Technical Structure:**
    (Briefly comment on abstract, keywords, and grammar in the detected language).
    
    **3. (Translate): Quality Rating:** [Score out of 100]
    
    **4. (Translate): Required Actionable Corrections:**
    (List corrections ONLY if the paper is poor. Otherwise, skip or mention minor tweaks).
    
    **5. (Translate): Board's Final Verdict:**
    * **PROCEED TO PUBLISH:** [YES / NO] (Translate YES/NO)
    * **BOARD DECISION:** [APPROVED / REVISE SPECIFIC SECTIONS / DECLINED] (Translate properly)
    * **Final Remarks:** (A formal 1-2 sentence human-written concluding remark in the detected language).
    
    Here is the Research Paper Text:
    {text}
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"API Error: {e}"

# 5. Function to Create Word Document (.docx) from the Report
def create_docx(report_text):
    doc = Document()
    doc.add_heading('MIJMR Editorial Board Evaluation Report', 0)
    
    # Clean basic markdown tags
    clean_text = report_text.replace("**", "").replace("*", "")
    
    # Add content to the document
    for line in clean_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
            
    # Save document to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# 6. Streamlit User Interface
st.sidebar.header("🎯 Evaluation Settings")
evaluation_standard = st.sidebar.selectbox(
    "Select Paper Standard Level:",
    ["Normal Standard (Local/Tamil Nadu College Level)", "High Standard (International/Foreign Level)"]
)

uploaded_file = st.file_uploader("📄 Upload Submitted Paper (PDF format only)", type=["pdf"])

if uploaded_file is not None:
    st.success("File uploaded successfully!")
    
    if st.button("🚀 Run Editorial Board Evaluation"):
        with st.spinner("⏳ Reading and processing the document... Please wait..."):
            paper_text = extract_text_from_pdf(uploaded_file)
            
            if "Error extracting text" in paper_text:
                st.error(paper_text)
            elif len(paper_text.strip()) < 50:
                st.warning("⚠️ Text extraction failed or insufficient data. Ensure the PDF is not a scanned image.")
            else:
                st.info("✅ Document verified. Generating board statement in the paper's language...")
                time.sleep(2)
                
                # Step B: Editorial Analysis
                report = analyze_paper_with_gemini(paper_text, api_key, evaluation_standard)
                
                # Show Report on Screen
                st.markdown("---")
                st.markdown(report)
                st.markdown("---")
                st.success("Evaluation Process Completed Successfully.")
                
                # Generate Word Doc and provide Download Button
                docx_buffer = create_docx(report)
                st.download_button(
                    label="📥 Download Official Board Report (Word .docx)",
                    data=docx_buffer,
                    file_name="MIJMR_Editorial_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
